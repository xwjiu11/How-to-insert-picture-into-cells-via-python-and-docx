from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def insert_picture(document,rows,cols,picture_path,picture_width,picture_height):
	tables = document.add_table(rows = rows,cols = cols)
	para = tables.rows[0].cells[0].add_paragraph()
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	r = para.add_run()
	r.add_picture('{picture_path}'.format(picture_path = picture_path),width = Inches(picture_width),height = Inches(picture_height))

